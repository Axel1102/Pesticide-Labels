from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_print_materials import (
    COUNTERBALANCE,
    GREEN,
    GOLD,
    INK,
    LIGHT,
    LINE,
    MUTED,
    OUTPUT_DIR,
    RED,
    add_info_strip,
    add_kicker,
    add_notice,
    add_page_header,
    add_section_gap,
    add_simple_table,
    add_title,
    clear_cell,
    page_break,
    set_cell_border,
    set_cell_margins,
    set_cell_shading,
    set_paragraph,
    set_run,
    set_table_geometry,
    style_doc,
)


HOST_DOCX = OUTPUT_DIR / "农药标签实验_主持人实验指南.docx"
PARTICIPANT_SCRIPT_DOCX = OUTPUT_DIR / "农药标签实验_参与者说明与宣读稿.docx"


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    set_run(p.add_run(text), size=13, bold=True, color="2C5F58")


def add_h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 3")
    set_run(p.add_run(text), size=11.5, bold=True, color="315C56")


def add_body(doc: Document, text: str, *, bold: bool = False, color: str = INK, after: float = 5) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, after=after, line=1.24)
    set_run(p.add_run(text), size=10.5, bold=bold, color=color)


def add_bullets(doc: Document, items: list[str]) -> None:
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph(p, after=4, line=1.22)
        set_run(p.add_run(text), size=10.3, color=INK)


def add_steps(doc: Document, items: list[str]) -> None:
    """Compact task-local numbered steps that always restart from 1."""
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [12, 170])
    for index, text in enumerate(items, start=1):
        cells = table.add_row().cells
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell, color="FFFFFF", size=0)
            set_cell_margins(cell, top=65, bottom=65, start=90, end=90)
        set_cell_shading(cells[0], "EEF3F1")
        p0 = clear_cell(cells[0])
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p0.add_run(str(index)), size=9.2, bold=True, color="2C5F58")
        p1 = clear_cell(cells[1])
        set_paragraph(p1, after=0, line=1.16)
        set_run(p1.add_run(text), size=9.8, color=INK)
    set_table_geometry(table, [12, 170])
    add_section_gap(doc, 4)


def add_script_box(doc: Document, text: str, *, label: str = "逐字读取", fill: str = "F1F5F4") -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [182])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="2C5F58", size=10, sides=("left",))
    set_cell_margins(cell, top=150, bottom=150, start=190, end=170)
    p = clear_cell(cell)
    set_paragraph(p, after=4)
    set_run(p.add_run(label + "\n"), size=8.5, bold=True, color="2C5F58")
    set_run(p.add_run(text), size=10.5, color=INK)
    add_section_gap(doc, 5)


def add_checklist(doc: Document, items: list[str], *, font_size: float = 9.5) -> None:
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [10, 172])
    for index, text in enumerate(items):
        cells = table.add_row().cells
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell, color="D6DDDA", size=5)
            set_cell_margins(cell, top=95, bottom=95, start=110, end=110)
            set_cell_shading(cell, "FFFFFF" if index % 2 == 0 else "F8F9F8")
        p0 = clear_cell(cells[0])
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p0.add_run("□"), size=11, color="2C5F58")
        p1 = clear_cell(cells[1])
        set_paragraph(p1, after=0, line=1.16)
        set_run(p1.add_run(text), size=font_size, color=INK)
    set_table_geometry(table, [10, 172])
    add_section_gap(doc, 5)


def build_host_guide() -> None:
    doc = Document()
    style_doc(doc, "农药标签实验 · 主持人实验指南")

    add_kicker(doc, "Facilitator Guide", "2C5F58")
    add_title(doc, "农药标签实验主持人指南", "从会前准备到结束归档的标准执行流程 · 不向参与者展示")
    add_notice(doc, "仅限实验主持人使用。不得向参与者展示标准答案、评分边界或标签条件名称。所有产品与药液均为实验模拟材料。")
    add_info_strip(
        doc,
        [("研究结构", "3 个 Block / 人"), ("Block 顺序", "T1 → T2 → T3 → T4"), ("标签条件", "Original / Enhanced / Dynamic"), ("核心原则", "统一话术 · 不提示")],
        [45.5, 45.5, 45.5, 45.5],
    )
    add_section_gap(doc, 8)
    add_h2(doc, "本指南的使用范围")
    add_body(doc, "本指南负责统一实验实施，不替代知情同意文件、招募材料或研究者答案册。正式实验前，应先完成伦理审批要求、主持人培训、设备测试与小规模预试。")
    add_h2(doc, "主持人的四条底线")
    add_steps(
        doc,
        [
            "只读统一话术；可以原句重复，但不自行解释标签、术语、计算或题意。",
            "不确认答案是否正确，不指向标签位置，不替参与者点击动态标签或选择工具。",
            "同一环节对所有参与者使用相同材料、摆放方式、计时口径和追问规则。",
            "配药容器不设规定容量；记录选择但不以容器大小或桶数计分。",
        ],
    )
    add_notice(doc, "遇到设备故障、液体洒漏或参与者不适时，以安全和自愿参与为先：暂停计时、处理问题并记录中断。", fill="FFF8E9", color=GOLD)

    page_break(doc)
    add_page_header(doc, "BEFORE THE SESSION", "会前准备与材料核对", note="建议在首位参与者到达前完成一次完整走台。")
    add_h2(doc, "研究日开始前")
    add_checklist(
        doc,
        [
            "确认当日参与者编号清单、G1-G6 分组顺序和预约时间；答题材料只写编号，不写姓名。",
            "按 Block 分装三个材料袋：正确标签版本、对应产品 T1 页、T2 四张卡、T3 提示卡、T4 情境卡、两页记录表。",
            "检查 Original 与 Enhanced 打印清楚；Dynamic 设备电量、网络或离线版本、触控与日志功能正常。",
            "将动态标签恢复到统一首页，关闭历史页面、搜索结果、弹窗和上一个参与者留下的状态。",
            "核对 T2 每套 4 张真实图片卡和 3 张分类区标题卡均齐全、编号与产品一致。",
            "核对 T3 模拟药瓶、有色水、清水、托盘、毛巾、搅拌工具、漏斗和所有量具；严禁放入真实农药。",
            "提供相同的一组可选容器，外观与摆放保持一致；不标记推荐容器，不暗示某个容量更正确。",
            "准备独立计时器；若方案包含录音或录像，只在完成相应同意后启用。",
            "把研究者答案册放在参与者看不到的位置；评分原则上在参与者离场后完成。",
        ],
    )
    add_h2(doc, "桌面标准布局")
    add_simple_table(
        doc,
        ["区域", "摆放内容", "控制要求"],
        [
            ("参与者正前方", "当前标签 + 当前任务材料", "一次只出现一个任务；上一任务材料及时收回。"),
            ("参与者侧前方", "计量工具与可选容器", "所有人数量、位置和朝向一致。"),
            ("主持人一侧", "记录表、计时器、备用材料", "答案册遮挡；不要在答题时打分。"),
            ("操作区", "防水托盘、清水、模拟药液、清洁用品", "保持干燥；每个 Block 后完全复位。"),
        ],
        [33, 63, 86],
        font_size=8.7,
    )

    page_break(doc)
    add_page_header(doc, "ASSIGNMENT", "分组、Block 顺序与签到")
    add_h2(doc, "G1-G6 分配表")
    add_simple_table(doc, ["组别", "Block 1", "Block 2", "Block 3"], [list(row) for row in COUNTERBALANCE], [25, 52, 52, 53], font_size=9.2)
    add_section_gap(doc, 7)
    add_h2(doc, "条件代码")
    add_info_strip(doc, [("O", "Original"), ("E", "Enhanced Static"), ("D", "Dynamic"), ("产品", "A / B / C")], [45.5] * 4)
    add_section_gap(doc, 7)
    add_h2(doc, "参与者到达后")
    add_checklist(
        doc,
        [
            "核对预约身份后，只在实验材料上填写参与者编号；按批准流程完成知情同意。",
            "若允许并已同意录音/录像，告知设备已启动；未同意则关闭相应记录。",
            "填写组别和三个 Block 顺序；再次核对每个产品与标签条件只出现一次。",
            "请参与者将真实农药使用经验作为背景，不把实验标签内容用于现实生产。",
            "发放参与者说明稿并逐字宣读；完成四项理解确认后才开始 Block 1。",
        ],
    )
    add_notice(doc, "不要把标签条件说成“普通版/更好版/智能版”，也不要评价任何版本。对参与者只说“当前标签”。", fill="FFF8E9", color=GOLD)

    page_break(doc)
    add_page_header(doc, "RUN OF SHOW", "单个 Block 的标准流程与计时")
    add_simple_table(
        doc,
        ["阶段", "主持人操作", "计时开始", "计时停止"],
        [
            ("标签呈现", "展示本 Block 指定标签；说明可全程查看。", "不单独计时", "-"),
            ("T1 判断", "发 5 题判断页；不逐题解释。", "读完 T1 指令并示意开始", "最后一题落笔/明确口答"),
            ("T2 分类", "摆好 3 个区域和 4 张卡。", "读完 T2 指令，手离开材料", "参与者说完成或停止移动"),
            ("T3 配药", "发任务卡；仅处理必要的容器交接。", "读完 T3 指令", "参与者说准备完成"),
            ("T4 调整", "紧接 T3 发情境卡。", "读完情境与问题", "给出最终方案和理由"),
            ("Block 重置", "收回全部材料并恢复桌面。", "不计入任务时间", "复位完成"),
        ],
        [24, 72, 43, 43],
        font_size=8.3,
    )
    add_section_gap(doc, 8)
    add_h2(doc, "统一 Block 开场")
    add_script_box(doc, "下面开始新一轮。你会看到一个产品和一份当前标签。接下来的四个任务都根据这份标签完成。整个过程中都可以继续查看标签。每道题中没有特别提到的情况，都可以认为符合正常使用条件。")
    add_h2(doc, "时间记录规则")
    add_bullets(
        doc,
        [
            "不催促、不报剩余时间；若正式方案设置上限，必须在预试后锁定，并对所有人一致。",
            "主持人造成的等待、设备修复或取放重物时间单独记为“实验中断”，不混入主动作答时间。",
            "参与者边操作边说话时继续计时；最终确认后停止，不因随后闲聊延长。",
            "漏记时间时不要估算，记录为缺失并说明原因。",
        ],
    )

    page_break(doc)
    add_page_header(doc, "STANDARD LANGUAGE", "允许的话术、禁止提示与求助处理")
    add_h2(doc, "允许原句使用的中性回应")
    neutral_rows = [
        ("要求重复题目", "我可以把刚才的原话再读一遍。"),
        ("询问答案是否正确", "我现在不能告诉你对错，请按你的判断完成。"),
        ("不知道怎么办", "请根据当前标签和桌面材料，按你认为合适的方式完成；不确定时可以明确说不确定。"),
        ("询问标签术语含义", "这个任务需要你按照自己对当前标签的理解判断，我不能补充解释。"),
        ("询问容器是否正确", "容器大小和数量没有规定，请自行选择；我不能评价你的选择。"),
        ("要求计算帮助", "我不能帮助计算。你可以继续查看标签，并用桌面上的工具完成。"),
        ("需要暂停", "可以，我们先暂停。准备好继续时请告诉我。"),
    ]
    add_simple_table(doc, ["参与者情况", "主持人统一回应"], neutral_rows, [55, 127], font_size=8.8)
    add_section_gap(doc, 7)
    add_h2(doc, "不得做的事")
    add_bullets(
        doc,
        [
            "指向、圈出或翻到标签中的答案位置；用语气、表情或停顿暗示正误。",
            "解释“有机磷”“含铜制剂”“碱性药剂”等类别，或替参与者判断图片卡。",
            "提示总水量、药量、最高剂量、计算公式、正确工具或操作顺序。",
            "为动态标签演示搜索路径、代替点击，或只在某一条件下主动介绍功能。",
            "在 T4 前评价 T3 方案，或提醒参与者注意某个情境限制。",
        ],
    )
    add_notice(doc, "每次求助都记录发生的任务、原始问题和主持人回应。技术故障与内容求助要分开编码。", fill="EEF6F0", color=GREEN)

    page_break(doc)
    add_page_header(doc, "T1 + T2", "判断与混配分类的主持流程")
    add_h2(doc, "T1 · 单变量判断")
    add_script_box(doc, "下面有 5 句话。请根据当前标签，逐题选择“对”“错”或“不确定”。每题只根据题目明确写出的情况判断；没有提到的情况都认为符合正常使用条件。完成后告诉我。")
    add_steps(
        doc,
        [
            "发放与当前产品对应的一页 T1，不要同时发其他产品题目。",
            "开始计时；观察标签回看，但不要站到参与者身后或盯住视线位置。",
            "参与者如漏题，只说“请确认每题都有一个选择”，不指出具体答案。",
            "停止计时并立即收回题页；不要当场核对或讨论。",
        ],
    )
    add_h2(doc, "T2 · 混配分类")
    add_script_box(doc, "现在有 4 张其他产品卡。请只根据当前标签，把它们分别放到“明确不能混用”“标签未发现明确禁忌”或“不确定”三个区域。可以移动和修改，确定后告诉我。")
    add_steps(
        doc,
        [
            "三个分类区从左到右固定摆放；四张卡在下方按卡号顺序正面朝上，图片、名称和有效成分均可见。",
            "读完指令后手离开材料并开始计时；参与者可以反复移动卡片。",
            "只记录最终位置；如研究计划记录移动过程，可另记首次放置与修改次数。",
            "停止计时、拍摄或抄录最终布局（仅在已获同意时），随后收回全部卡片并核对 4 张齐全。",
        ],
    )
    add_notice(doc, "“标签未发现明确禁忌”只是一种实验分类，不代表现实农业中一定可以混用。主持人不得把它解释成“安全”。", fill="FFF8E9", color=GOLD)

    page_break(doc)
    add_page_header(doc, "T3", "实际配药任务的主持流程", note="容器不规定；记录容器与桶数，但不按容器选择计分。")
    add_h2(doc, "标准提示")
    add_script_box(doc, "现在需要为 1 亩地配制一次喷施需要的药液。请根据当前标签，使用桌面上的东西完成配药。容器大小和数量没有规定，请你自己决定水量、药量、量取工具和配药顺序；可以分一桶或多桶完成。准备好后告诉我。")
    add_h2(doc, "标准操作顺序（主持人动作，不是参与者答案）")
    add_steps(
        doc,
        [
            "确认桌上只有实验模拟液，开始计时；不介绍工具用途，也不说任何容量是推荐值。",
            "让参与者自行选择容器、桶数和目标水量。主持人只记录，不判断。",
            "若采用“先选方案、再提供约半水容器”的减负设置，在参与者明确方案后暂停计时；容器交到位后恢复计时。所有参与者必须使用同一设置。",
            "观察并记录参与者实际或口述的每桶水量、药量、量具、计算尝试、标签回看、混合步骤和求助。",
            "参与者说完成后停止计时，先记录最终方案，再进行 T4；此时不要评价 T3。",
        ],
    )
    add_h2(doc, "安全与异常")
    add_simple_table(
        doc,
        ["情况", "处理"],
        [
            ("液体洒出", "立即暂停；请参与者后退，主持人用托盘/毛巾处理并记录中断。"),
            ("接触眼口或明显不适", "立即停止该任务并按实验场地安全预案处理。"),
            ("参与者想品尝/闻液体", "立即制止并说明材料不得入口或靠近面部。"),
            ("容器过重", "由主持人协助搬动，但不替参与者决定水量或配方；记录协助。"),
            ("材料用尽或量具损坏", "暂停计时、更换同规格材料，记录设备中断后继续。"),
        ],
        [50, 132],
        font_size=8.7,
    )
    add_notice(doc, "正式实验前必须锁定“全水量操作”或“统一减负模拟”中的一种，不得因参与者不同临时改变。", fill="FFF2EF", color=RED)

    page_break(doc)
    add_page_header(doc, "T4 + DYNAMIC", "情境调整、动态标签与偏差处理")
    add_h2(doc, "T4 · 情境调整")
    add_script_box(doc, "刚才是一般情况下的配药。现在再告诉你一个情况。请看这张卡。知道这个情况以后，你这一次准备怎么用？请把最终方案和理由说完整。")
    add_steps(
        doc,
        [
            "保持同一标签可见，T3 材料与参与者最终方案不做评价；发放当前产品的 T4 情境卡。",
            "逐字读卡片和问题，开始计时。若回答不完整，只能追问一次：“请把最终方案说完整，包括是否改变、准备怎么用以及理由。”",
            "记录是否继续使用、最终浓度或总量、调整方向、理由、标签回看和信心；不在现场评分。",
            "参与者确认最终方案后停止计时，收回材料并结束当前 Block。",
        ],
    )
    add_h2(doc, "Dynamic 条件的额外控制")
    add_bullets(
        doc,
        [
            "每个 Dynamic Block 从相同首页和默认状态开始；主持人不主动介绍页面结构或可用功能。",
            "内容求助按统一中性话术回应；只有无法点击、页面崩溃、设备断电等技术故障可以介入。",
            "技术介入时暂停计时，记录故障类型、持续时间和主持人动作；恢复到故障前可比状态后继续。",
            "Block 结束后导出所需日志（若已批准并告知），再清除参与者输入和浏览状态。",
        ],
    )
    add_h2(doc, "实验偏差与停止规则")
    add_simple_table(
        doc,
        ["事件", "是否继续", "记录方式"],
        [
            ("发错产品/标签/卡片", "停止该 Block；不要用错误材料继续", "记重大偏差并联系负责人决定重做或剔除"),
            ("主持人无意提示答案", "完成当前任务后评估", "逐字记录提示内容、时间点和受影响题目"),
            ("参与者主动中止", "立即停止", "不追问原因；按同意流程处理已收集数据"),
            ("短暂休息", "可继续", "记录开始/结束；休息安排在任务之间"),
            ("安全事件或持续不适", "立即停止", "按安全与伦理流程上报"),
        ],
        [52, 52, 78],
        font_size=8.4,
    )

    page_break(doc)
    add_page_header(doc, "CLOSEOUT", "结束、复位与一页式检查")
    add_h2(doc, "实验结束话术")
    add_script_box(doc, "今天的任务已经完成。这些产品、剂量和药液都是为实验制作的模拟材料，不能用于现实农业生产。我们关注的是不同标签怎样帮助人们找到并使用信息，不是在评价你的农业知识或个人能力。现在还有没有想补充的感受或问题？")
    add_h2(doc, "参与者离场后")
    add_checklist(
        doc,
        [
            "停止录音/录像并确认文件保存；按参与者编号命名，不在任务文件名中写姓名。",
            "核对 3 个 Block 的 T1-T4 记录、时间、求助、标签回看和偏差说明是否完整。",
            "在参与者离场后依据研究者手册评分；无法判断的项目标记待复核，不凭印象补填。",
            "回收全部 T2 卡、T3/T4 卡和标签；清洗量具、补充模拟液并擦干操作区。",
            "将 Dynamic 标签恢复首页并清除上一位参与者状态；检查下一个分组材料袋。",
            "纸质材料放入编号文件夹；电子文件按批准的数据管理方案保存。",
        ],
        font_size=9.2,
    )
    add_h2(doc, "开场前 30 秒检查")
    add_info_strip(doc, [("编号", "已填"), ("组别", "已核"), ("材料", "3 个 Block 齐全"), ("设备", "计时/动态/记录正常")], [45.5] * 4)
    add_section_gap(doc, 6)
    add_h2(doc, "每个 Block 结束 30 秒检查")
    add_info_strip(doc, [("T1", "5 题完整"), ("T2", "4 卡位置已记"), ("T3", "最终量与步骤已记"), ("T4", "方案+理由已记")], [45.5] * 4)
    add_section_gap(doc, 7)
    doc.save(HOST_DOCX)


def build_participant_script() -> None:
    doc = Document()
    style_doc(doc, "农药标签实验 · 参与者说明与宣读稿")

    add_kicker(doc, "Participant Information", "2C5F58")
    add_title(doc, "参与者实验说明", "主持人逐字宣读版 + 参与者桌面简版")
    add_notice(doc, "实验中的产品、标签信息和药液均为模拟材料，仅用于研究，不能作为现实农业用药依据。")
    add_info_strip(doc, [("参与者编号", "________________"), ("日期", "____年__月__日"), ("预计结构", "3 轮 × 4 个任务")], [70, 55, 57])
    add_section_gap(doc, 8)
    add_h2(doc, "这项实验在做什么")
    add_body(doc, "我们想了解人们在不同形式的农药标签上，怎样寻找信息、理解限制并完成判断和模拟操作。我们关注的是标签设计是否容易使用，不是在评价你的农业知识、文化程度、阅读能力或计算能力。")
    add_h2(doc, "参加时需要知道")
    add_bullets(
        doc,
        [
            "你会连续完成 3 轮任务。每轮使用一个产品和一份当前标签，并依次完成判断、混配分类、模拟配药和情境调整。",
            "整个任务中都可以查看当前标签；会记录答案、完成时间、操作过程、标签回看和求助情况。",
            "如果本次研究包含录音或录像，只会按你已经同意的方式启用。",
            "你可以提出暂停或停止；如有不适、液体洒出或不想继续，请立即告诉主持人。",
            "实验材料不得入口、不得靠近眼睛和口鼻；请勿把实验标签带回现实生产中使用。",
        ],
    )
    add_h2(doc, "统一判断原则")
    add_script_box(doc, "每道题中没有特别提到的情况，都可以认为符合正常使用条件。请只根据题目给出的情况和当前标签判断。不能确定时，可以选择或明确说“不确定”。", label="请特别记住")
    add_notice(doc, "主持人可以重复原话和处理设备问题，但不能告诉你答案、解释标签术语、帮助计算或提示在哪里找信息。", fill="FFF8E9", color=GOLD)

    page_break(doc)
    add_page_header(doc, "READ ALOUD", "开始实验前的统一宣读稿", note="主持人按原文完整读一遍；参与者可以同时查看本页。")
    add_script_box(
        doc,
        "欢迎参加这项实验。今天使用的产品、标签内容和药液都是研究用的模拟材料，不是真实农业建议。我们想了解不同标签怎样帮助人们找到和使用信息，不是在考你的农业知识。\n\n"
        "你会完成三轮任务。每一轮都使用一个产品和一份当前标签，并依次完成四个任务。整个过程中，你都可以继续查看当前标签。每道题中没有特别提到的情况，都可以认为符合正常使用条件。请只根据题目和当前标签判断。\n\n"
        "如果不能确定，可以选择或告诉我“不确定”。我可以重复刚才的原话，也可以处理设备故障，但不能告诉你答案、解释标签中的词、帮你计算，或提示应该看哪里。请按照你自己的理解完成。\n\n"
        "模拟配药时，容器大小和数量没有规定，你可以自行选择一桶或多桶，并自己决定水量、药量、量取工具和操作顺序。桌上的液体不能入口，也不要接触眼睛和口鼻；如果洒出或感到不适，请马上停下并告诉我。\n\n"
        "实验会记录你的答案、用时和操作过程。你可以随时要求暂停或停止。准备好以后，我们开始第一轮。",
        label="主持人逐字读取",
        fill="F4F7F6",
    )
    add_h2(doc, "开始前理解确认")
    add_simple_table(
        doc,
        ["主持人确认问题", "参与者确认"],
        [
            ("完成任务时可以一直查看当前标签，对吗？", "□ 是  □ 需再读一遍"),
            ("题目没提到的情况，按正常条件理解，对吗？", "□ 是  □ 需再读一遍"),
            ("不知道时可以选择“不确定”，主持人不会提供答案，对吗？", "□ 是  □ 需再读一遍"),
            ("配药容器大小和数量没有规定，由你自行选择，对吗？", "□ 是  □ 需再读一遍"),
        ],
        [125, 57],
        font_size=9,
    )
    add_section_gap(doc, 7)
    add_notice(doc, "如有一项未确认，主持人只重复对应的原句，不增加例子或解释。", fill="EEF6F0", color=GREEN)

    page_break(doc)
    add_page_header(doc, "QUICK REFERENCE", "四个任务的桌面简版说明", note="可在实验过程中放在桌边；不包含任何产品答案或剂量。")
    add_simple_table(
        doc,
        ["任务", "你需要做什么", "完成时告诉主持人"],
        [
            ("T1 判断", "根据当前标签，对 5 句话选择“对 / 错 / 不确定”。", "全部 5 题已作答"),
            ("T2 混配分类", "把 4 张产品卡放入三个分类区；可以修改位置。", "所有卡已放好"),
            ("T3 模拟配药", "为 1 亩地完成一次模拟配药；自行决定容器、水量、药量、量具和顺序。", "最终方案和操作已完成"),
            ("T4 情境调整", "阅读新增情况，说出最终准备怎么用以及为什么。", "最终方案和理由已说完整"),
        ],
        [28, 105, 49],
        font_size=8.8,
    )
    add_section_gap(doc, 8)
    add_h2(doc, "你可以")
    add_bullets(
        doc,
        [
            "随时查看当前标签。",
            "修改尚未提交的答案或卡片位置。",
            "说出自己的计算过程和判断理由。",
            "选择或说“不确定”。",
            "要求主持人重复原话、暂停或停止。",
        ],
    )
    add_h2(doc, "请不要")
    add_bullets(
        doc,
        [
            "向主持人询问哪个答案正确、某个术语是什么意思或应该用哪个工具。",
            "品尝、闻嗅或让模拟液接触眼睛、口鼻；液体洒出时请立即停下。",
            "把实验中的产品、剂量、混配结论或操作方法用于现实农业生产。",
        ],
    )
    add_notice(doc, "有任何不清楚或不舒服的地方，请直接告诉主持人。", fill="F7F8F7", color="2C5F58")

    doc.save(PARTICIPANT_SCRIPT_DOCX)


if __name__ == "__main__":
    build_host_guide()
    build_participant_script()
    print(HOST_DOCX)
    print(PARTICIPANT_SCRIPT_DOCX)
