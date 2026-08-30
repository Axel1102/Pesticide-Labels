from __future__ import annotations

import copy
import math
import os
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(r"C:\Users\57256\Documents\ChatGPT\农药")
IMAGE_DIR = ROOT / "dynamic-label-demo" / "public" / "images"
OUTPUT_DIR = ROOT / "output" / "print_materials"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

PARTICIPANT_DOCX = OUTPUT_DIR / "农药标签实验_参与者打印材料.docx"
RESEARCHER_DOCX = OUTPUT_DIR / "农药标签实验_研究者手册与记录表.docx"

FONT = "Microsoft YaHei"
MONO = "Consolas"
INK = "1E2523"
MUTED = "65706C"
LIGHT = "F3F5F4"
LINE = "B9C2BE"
RED = "A3372A"
GOLD = "9A6B12"
GREEN = "287545"
PRODUCT_COLORS = {"A": "2C708D", "B": "377748", "C": "8B5A2B"}
PRODUCT_LIGHT = {"A": "EAF3F7", "B": "EDF5EF", "C": "F6F0EA"}


T1 = {
    "A": [
        ("A1", "玉米能数出4片展开的叶子时，可以用这款药。", "对", "玉米3-5叶期可用。"),
        ("A2", "玉米能数出6片展开的叶子时，还可以用这款药。", "错", "超过玉米5叶期的使用上限。"),
        ("A3", "这个季节用过一次这款药，过一周还可以再用一次。", "错", "每季最多使用1次；间隔一周也不能突破次数上限。"),
        ("A4", "喷药时，药液全面喷洒在地里覆盖所有的植物。", "错", "应对杂草茎叶均匀喷雾，并避开玉米心叶。"),
        ("A5", "天气预报说大约半小时后会下雨，今天可以照常喷。", "错", "预计1小时内降雨时不得施药。"),
    ],
    "B": [
        ("B1", "小麦叶子刚开始出现少量白色粉斑时，可以开始用这款药。", "对", "发病前或发病初期可开始施药。"),
        ("B2", "5天前刚喷过一次这款药，今天可以再喷一次。", "错", "再次使用至少间隔7天。"),
        ("B3", "这个季节已经喷过两次这款药，还可以再喷一次。", "错", "每季最多使用2次。"),
        ("B4", "还有18天就准备收麦，这时可以用这款药。", "错", "收获前20天内不得使用。"),
        ("B5", "使用这款药时，要把药液均匀喷在小麦的茎和叶上。", "对", "标签要求对小麦茎、叶均匀周到喷雾。"),
    ],
    "C": [
        ("C1", "虫子还比较小，主要藏在玉米心叶里的时候，可以用这款药。", "对", "卵孵化盛期至低龄幼虫期可用。"),
        ("C2", "4天前喷过一次这款药，今天可以再喷一次。", "错", "再次使用至少间隔7天。"),
        ("C3", "这个季节已经喷过三次这款药，还可以再喷一次。", "错", "每季最多使用3次。"),
        ("C4", "还有10天准备收玉米，这时还可以用这款药。", "错", "收获前14天内不得使用。"),
        ("C5", "喷这款药时，要重点喷玉米心叶和有虫害的地方。", "对", "标签要求重点喷心叶和受害部位。"),
    ],
}


MIXING = {
    "A": [
        ("A-T2-01", "辛硫磷乳油", "有效成分：40% 辛硫磷", "phoxim-product-photo-cutout.png", "明确不能混用"),
        ("A-T2-02", "敌敌畏乳油", "有效成分：80% 敌敌畏", "dichlorvos-product-photo-cutout.png", "明确不能混用"),
        ("A-T2-03", "氢氧化铜可湿性粉剂", "有效成分：77% 氢氧化铜", "copper-hydroxide-product-real.jpg", "标签未发现明确禁忌"),
        ("A-T2-04", "枯草芽孢杆菌制剂", "有效成分：枯草芽孢杆菌", "biofungicide-product-real.png", "标签未发现明确禁忌"),
    ],
    "B": [
        ("B-T2-01", "波尔多液可湿性粉剂", "有效成分：80% 波尔多液", "bordeaux-mixture-product-real.jpg", "明确不能混用"),
        ("B-T2-02", "氢氧化铜可湿性粉剂", "有效成分：77% 氢氧化铜", "copper-hydroxide-product-real.jpg", "明确不能混用"),
        ("B-T2-03", "辛硫磷乳油", "有效成分：40% 辛硫磷", "phoxim-product-photo-cutout.png", "标签未发现明确禁忌"),
        ("B-T2-04", "矿物油助剂", "有效成分：农业用矿物油", "mineral-oil-product-real.jpg", "标签未发现明确禁忌"),
    ],
    "C": [
        ("C-T2-01", "石硫合剂水剂", "有效成分：29% 石硫合剂", "lime-sulfur-product-real.jpg", "明确不能混用"),
        ("C-T2-02", "波尔多液可湿性粉剂", "有效成分：80% 波尔多液", "bordeaux-mixture-product-real.jpg", "明确不能混用"),
        ("C-T2-03", "枯草芽孢杆菌制剂", "有效成分：枯草芽孢杆菌", "biofungicide-product-real.png", "标签未发现明确禁忌"),
        ("C-T2-04", "辛硫磷乳油", "有效成分：40% 辛硫磷", "phoxim-product-photo-cutout.png", "标签未发现明确禁忌"),
    ],
}


T4 = {
    "A": "田里的杂草明显比一般情况密，但目前仍然比较矮小，也还在适合施药的时期。这个季节还没有使用过这款药。\n\n这种情况下，你这一次准备怎么用？",
    "B": "上一次使用 Product B 是8天前，这个季节只使用过1次，现在仍然处于适合使用的时期。但是上一次使用以后感觉效果明显比以前差。\n\n这一次你准备怎么用？",
    "C": "现在看到的仍然主要是低龄幼虫，但是虫量比一般情况下多。上一次使用后的效果也不太理想，目前距离上次施药已经超过7天。\n\n这一次你准备怎么用？",
}


COUNTERBALANCE = [
    ("G1", "A-Original", "B-Enhanced", "C-Dynamic"),
    ("G2", "B-Dynamic", "C-Original", "A-Enhanced"),
    ("G3", "C-Enhanced", "A-Dynamic", "B-Original"),
    ("G4", "C-Dynamic", "B-Enhanced", "A-Original"),
    ("G5", "A-Enhanced", "C-Original", "B-Dynamic"),
    ("G6", "B-Original", "A-Dynamic", "C-Enhanced"),
]


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run(run, size=10.5, bold=False, color=INK, font=FONT, italic=False):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = rgb(color)
    return run


def set_paragraph(p, before=0, after=5, line=1.2, align=None, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_together = keep
    if align is not None:
        p.alignment = align
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LINE, size=8, sides=("top", "left", "bottom", "right"), style="single"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in sides:
        tag = f"w:{side}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), style)
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_mm, indent_dxa=0):
    total = sum(widths_mm)
    widths_dxa = [round(mm / 25.4 * 1440) for mm in widths_mm]
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(widths_dxa):
                continue
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Mm(widths_mm[i])


def hide_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        borders.append(el)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def clear_cell(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph(p, after=0)
    return p


def style_doc(doc: Document, running_title: str):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(13)
    section.bottom_margin = Mm(13)
    section.left_margin = Mm(14)
    section.right_margin = Mm(14)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color, before, after in (
        ("Title", 25, INK, 0, 5),
        ("Heading 1", 17, "2C5F58", 12, 7),
        ("Heading 2", 13, "2C5F58", 9, 5),
        ("Heading 3", 11.5, "315C56", 7, 4),
    ):
        st = doc.styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(size)
        st.font.color.rgb = rgb(color)
        st.font.bold = name != "Title"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for style_name in ("List Number", "List Bullet"):
        st = doc.styles[style_name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.2

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph(hp, after=0)
    set_run(hp.add_run(running_title), size=8.5, bold=True, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(fp, after=0)
    set_run(fp.add_run("实验用虚构材料 · 禁止用于真实农业生产    "), size=8, color=MUTED)
    run = fp.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run(run, size=8, color=MUTED)


def add_kicker(doc, text, color="2C5F58"):
    p = doc.add_paragraph()
    set_paragraph(p, after=2)
    set_run(p.add_run(text.upper()), size=9, bold=True, color=color)
    return p


def add_title(doc, title, subtitle=None, product=None):
    color = PRODUCT_COLORS.get(product, "2C5F58")
    p = doc.add_paragraph(style="Title")
    set_paragraph(p, after=4, keep=True)
    set_run(p.add_run(title), size=24, bold=True, color=INK)
    if subtitle:
        sub = doc.add_paragraph()
        set_paragraph(sub, after=9)
        set_run(sub.add_run(subtitle), size=11, color=MUTED)
    bar = doc.add_table(rows=1, cols=1)
    set_table_geometry(bar, [182])
    hide_table_borders(bar)
    set_cell_shading(bar.cell(0, 0), color)
    set_cell_margins(bar.cell(0, 0), top=45, bottom=45, start=0, end=0)
    clear_cell(bar.cell(0, 0))
    spacer = doc.add_paragraph()
    set_paragraph(spacer, after=5)


def add_page_header(doc, task, title, product=None, note=None):
    color = PRODUCT_COLORS.get(product, "2C5F58")
    add_kicker(doc, task, color)
    p = doc.add_paragraph()
    set_paragraph(p, after=2, keep=True)
    set_run(p.add_run(title), size=19, bold=True, color=INK)
    if note:
        p2 = doc.add_paragraph()
        set_paragraph(p2, after=7)
        set_run(p2.add_run(note), size=9.5, color=MUTED)
    bar = doc.add_table(rows=1, cols=1)
    set_table_geometry(bar, [182])
    hide_table_borders(bar)
    set_cell_shading(bar.cell(0, 0), color)
    set_cell_margins(bar.cell(0, 0), top=34, bottom=34, start=0, end=0)
    clear_cell(bar.cell(0, 0))
    p3 = doc.add_paragraph()
    set_paragraph(p3, after=2)


def add_notice(doc, text, fill="FFF2EF", color=RED):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [182])
    set_cell_shading(table.cell(0, 0), fill)
    set_cell_border(table.cell(0, 0), color=color, size=10, sides=("left",))
    set_cell_margins(table.cell(0, 0), top=100, bottom=100, start=160, end=140)
    p = clear_cell(table.cell(0, 0))
    set_run(p.add_run(text), size=9.5, bold=True, color=color)
    after = doc.add_paragraph()
    set_paragraph(after, after=3)


def add_info_strip(doc, items, widths=None, fill=LIGHT, accent="2C5F58"):
    if widths is None:
        widths = [182 / len(items)] * len(items)
    table = doc.add_table(rows=1, cols=len(items))
    set_table_geometry(table, widths)
    for i, (label, value) in enumerate(items):
        c = table.cell(0, i)
        set_cell_shading(c, fill)
        set_cell_border(c, color="D3D9D6", size=5)
        set_cell_margins(c, top=120, bottom=120, start=140, end=140)
        p = clear_cell(c)
        set_run(p.add_run(label + "\n"), size=8.3, bold=True, color=accent)
        set_run(p.add_run(value), size=10, bold=True, color=INK)
    return table


def add_question_block(doc, item, statement, product):
    color = PRODUCT_COLORS[product]
    fill = PRODUCT_LIGHT[product]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [18, 164])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in table.rows[0].cells:
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(c, top=145, bottom=145, start=160, end=160)
        set_cell_border(c, color=color, size=7)
    set_cell_shading(table.cell(0, 0), color)
    p0 = clear_cell(table.cell(0, 0))
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p0.add_run(item), size=12, bold=True, color="FFFFFF")
    set_cell_shading(table.cell(0, 1), fill)
    p1 = clear_cell(table.cell(0, 1))
    set_paragraph(p1, after=5, line=1.25)
    set_run(p1.add_run(statement), size=12, bold=True, color=INK)
    p2 = table.cell(0, 1).add_paragraph()
    set_paragraph(p2, after=0)
    set_run(p2.add_run("□ 对       □ 错       □ 不确定"), size=10.5, color=INK)
    gap = doc.add_paragraph()
    set_paragraph(gap, after=2)


def add_image_fitted(paragraph, path: Path, max_width_mm: float, max_height_mm: float, alt: str):
    with Image.open(path) as im:
        w, h = im.size
    scale = min(max_width_mm / w, max_height_mm / h)
    width = max(1, w * scale)
    height = max(1, h * scale)
    run = paragraph.add_run()
    drawing = run.add_picture(str(path), width=Mm(width), height=Mm(height))
    doc_pr = drawing._inline.docPr
    doc_pr.set("descr", alt)
    return drawing


def add_mixing_card(cell, card, product):
    card_id, name, ingredient, image_name, _ = card
    color = PRODUCT_COLORS[product]
    set_cell_border(cell, color=color, size=9)
    set_cell_margins(cell, top=110, bottom=110, start=130, end=130)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(f"与 PRODUCT {product} 混配"), size=8.2, bold=True, color=color)
    img_p = cell.add_paragraph()
    set_paragraph(img_p, after=3, align=WD_ALIGN_PARAGRAPH.CENTER, keep=True)
    add_image_fitted(img_p, IMAGE_DIR / image_name, 70, 61, name + "真实包装照片")
    name_p = cell.add_paragraph()
    set_paragraph(name_p, after=2, align=WD_ALIGN_PARAGRAPH.CENTER, keep=True)
    set_run(name_p.add_run(name), size=12, bold=True, color=INK)
    ing_p = cell.add_paragraph()
    set_paragraph(ing_p, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(ing_p.add_run(ingredient), size=9, color=MUTED)
    id_p = cell.add_paragraph()
    set_paragraph(id_p, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_run(id_p.add_run(card_id), size=7.5, color=MUTED, font=MONO)


def add_sort_zone(doc, title, instruction, fill, accent):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [182])
    c = table.cell(0, 0)
    set_cell_shading(c, fill)
    set_cell_border(c, color=accent, size=12)
    set_cell_margins(c, top=240, bottom=240, start=220, end=220)
    p = clear_cell(c)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(title), size=19, bold=True, color=accent)
    p2 = c.add_paragraph()
    set_paragraph(p2, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(p2.add_run(instruction), size=9.5, color=INK)
    gap = doc.add_paragraph()
    set_paragraph(gap, after=7)


def add_task_card(cell, product, title, body, task_code):
    color = PRODUCT_COLORS[product]
    fill = PRODUCT_LIGHT[product]
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=color, size=10)
    set_cell_margins(cell, top=210, bottom=210, start=220, end=220)
    p = clear_cell(cell)
    set_run(p.add_run(f"PRODUCT {product}"), size=9, bold=True, color=color)
    p2 = cell.add_paragraph()
    set_paragraph(p2, after=6, keep=True)
    set_run(p2.add_run(title), size=15, bold=True, color=INK)
    for i, para_text in enumerate(body.split("\n\n")):
        p3 = cell.add_paragraph()
        set_paragraph(p3, after=6 if i < len(body.split("\n\n")) - 1 else 4, line=1.3)
        set_run(p3.add_run(para_text), size=11, bold=(i == len(body.split("\n\n")) - 1), color=INK)
    p4 = cell.add_paragraph()
    set_paragraph(p4, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_run(p4.add_run(task_code), size=7.5, color=MUTED, font=MONO)


def add_simple_table(doc, headers, rows, widths_mm, font_size=9, header_fill="33413D"):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_mm)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        c = hdr.cells[i]
        set_cell_shading(c, header_fill)
        set_cell_border(c, color="FFFFFF", size=4)
        set_cell_margins(c, top=100, bottom=100, start=120, end=120)
        p = clear_cell(c)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(text), size=font_size, bold=True, color="FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            c = cells[i]
            set_cell_border(c, color=LINE, size=5)
            set_cell_margins(c, top=90, bottom=90, start=110, end=110)
            p = clear_cell(c)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(str(value)), size=font_size, color=INK)
    set_table_geometry(table, widths_mm)
    return table


def add_section_gap(doc, points=4):
    p = doc.add_paragraph()
    set_paragraph(p, after=points)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    set_paragraph(p, after=0)


def build_participant_doc():
    doc = Document()
    style_doc(doc, "农药标签实验 · 参与者打印材料")

    add_kicker(doc, "Participant Print Pack", "2C5F58")
    add_title(doc, "农药标签实验任务材料", "T1-T4 · 参与者版 · 按产品与任务分发")
    add_info_strip(doc, [("任务一", "判断现在是否符合"), ("任务二", "识别现实混配对象"), ("任务三", "完成一亩地配药"), ("任务四", "根据变化调整方案")], [45.5] * 4)
    add_section_gap(doc, 7)
    add_notice(doc, "所有产品、成分、登记信息、剂量和混配信息均为虚构，仅用于标签实验，禁止用于真实农业生产。")
    p = doc.add_paragraph(style="Heading 2")
    set_run(p.add_run("统一说明"), size=13, bold=True, color="2C5F58")
    instructions = [
        "每道题中没有特别提到的情况，都可以认为符合正常使用条件。只根据题目给出的情况和当前标签判断。",
        "可以在完成任务的过程中继续查看当前标签。",
        "如无法确定，可以选择“不确定”；不要向研究者询问答案。",
        "混配卡只用于分类判断，“标签未发现明确禁忌”不代表真实生产中一定可以混用。",
    ]
    for text in instructions:
        p = doc.add_paragraph(style="List Number")
        set_run(p.add_run(text), size=11, color=INK)
    add_section_gap(doc, 8)
    add_info_strip(doc, [("参与者编号", "________________"), ("组别", "G____"), ("日期", "____年__月__日")], [70, 45, 67])
    p = doc.add_paragraph()
    set_paragraph(p, before=14, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(p.add_run("研究者提示：打印后按产品与任务裁切/分发，勿一次性交给参与者。"), size=9, color=MUTED, italic=True)

    for product, product_name in (("A", "玉米除草剂"), ("B", "小麦杀菌剂"), ("C", "玉米杀虫剂")):
        page_break(doc)
        add_page_header(doc, "T1 · 单变量判断", f"Product {product} · {product_name}", product, "请逐题选择“对 / 错 / 不确定”。每题只根据题目写出的一个条件判断。")
        add_notice(doc, "没有特别提到的情况，都可以认为符合正常使用条件。", fill="F7F8F7", color=PRODUCT_COLORS[product])
        for item, statement, _, _ in T1[product]:
            add_question_block(doc, item, statement, product)
        p = doc.add_paragraph()
        set_paragraph(p, before=4, after=0)
        set_run(p.add_run("本页开始时间：__________    完成时间：__________    总体把握：□ 很有把握  □ 比较有把握  □ 不太确定"), size=9, color=MUTED)

    page_break(doc)
    add_page_header(doc, "T2 · 混配分类", "分类区标题卡", note="沿分隔处裁切。将4张药物卡放在相应标题卡附近。")
    add_sort_zone(doc, "明确不能混用", "当前标签明确禁止与这一产品混用", "FFF0ED", RED)
    add_sort_zone(doc, "标签未发现明确禁忌", "只表示当前标签没有明确写出，不代表现实中一定可以混用", "EEF6F0", GREEN)
    add_sort_zone(doc, "不确定", "根据当前标签仍无法判断", "FFF7E8", GOLD)

    for product, product_name in (("A", "玉米除草剂"), ("B", "小麦杀菌剂"), ("C", "玉米杀虫剂")):
        page_break(doc)
        add_page_header(doc, "T2 · 真实产品照片分类", f"Product {product} · {product_name}", product, "沿外框裁成4张卡。请根据名称、有效成分和包装照片判断，不要只看颜色或瓶形。")
        table = doc.add_table(rows=2, cols=2)
        set_table_geometry(table, [89, 89])
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, card in enumerate(MIXING[product]):
            add_mixing_card(table.cell(idx // 2, idx % 2), card, product)
        add_section_gap(doc, 2)
        p = doc.add_paragraph()
        set_paragraph(p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_run(p.add_run("照片复用自本实验增强静态/动态标签素材；实验中以卡片文字与当前标签为准。"), size=8, color=MUTED)

    page_break(doc)
    add_page_header(doc, "T3 · 实际配药", "任务提示卡", note="沿外框裁成3张卡；每个 block 只展示对应产品的一张。")
    table = doc.add_table(rows=3, cols=1)
    set_table_geometry(table, [182])
    prompt = "现在需要为1亩地配制一次喷施需要的药液。请根据这个标签，使用桌面上的东西完成配药。\n\n容器大小不作规定。你需要自己决定水量、药量、量取工具和配药顺序；可以分一桶或多桶完成。准备好后告诉研究者。"
    for i, product in enumerate(("A", "B", "C")):
        add_task_card(table.cell(i, 0), product, "请完成一次配药", prompt, f"{product}-T3")

    page_break(doc)
    add_page_header(doc, "T4 · 情境调整", "条件变化卡", note="沿外框裁成3张卡；在 T3 一般配药任务之后展示。")
    table = doc.add_table(rows=3, cols=1)
    set_table_geometry(table, [182])
    for i, product in enumerate(("A", "B", "C")):
        add_task_card(table.cell(i, 0), product, "现在再告诉你一个情况", T4[product], f"{product}-T4")

    doc.save(PARTICIPANT_DOCX)


def build_researcher_doc():
    doc = Document()
    style_doc(doc, "农药标签实验 · 研究者手册与记录表")

    add_kicker(doc, "Researcher Manual", "2C5F58")
    add_title(doc, "农药标签实验研究者手册", "流程、标准答案、配药标准、评分规则与可复印记录表")
    add_notice(doc, "本文件不得向参与者展示。所有答案只适用于本实验中的虚构产品与标签，不构成真实农业建议。")
    add_info_strip(doc, [("Block 结构", "一个产品 × 一个标签条件"), ("固定顺序", "T1 → T2 → T3 → T4"), ("核心控制", "每个任务只制造一种主要困难")], [62, 58, 62])
    p = doc.add_paragraph(style="Heading 2")
    set_run(p.add_run("每个 Block 的标准流程"), size=13, bold=True, color="2C5F58")
    steps = [
        "展示分配到的标签版本；确认参与者知道可以全程查看标签。",
        "T1：发放对应产品的5题判断页；记录总时长、答案和信心。",
        "T2：摆放3张分类区标题卡和对应产品4张混配卡；记录每张卡的最终位置。",
        "T3：发放对应产品的实际配药提示卡；观察容器、水量、药量、量具和操作顺序。",
        "T4：紧接 T3 展示对应情境卡；记录最终用量、调整理由和0-2分编码。",
        "收回本 block 全部材料，清理桌面，再开始下一 block。",
    ]
    for text in steps:
        p = doc.add_paragraph(style="List Number")
        set_run(p.add_run(text), size=10.5, color=INK)
    p = doc.add_paragraph(style="Heading 2")
    set_run(p.add_run("统一口头说明（逐字读取）"), size=13, bold=True, color="2C5F58")
    quote = doc.add_table(rows=1, cols=1)
    set_table_geometry(quote, [182])
    set_cell_shading(quote.cell(0, 0), "F1F5F4")
    set_cell_border(quote.cell(0, 0), color="2C5F58", size=9, sides=("left",))
    set_cell_margins(quote.cell(0, 0), top=130, bottom=130, start=180, end=160)
    q = clear_cell(quote.cell(0, 0))
    set_run(q.add_run("每道题中没有特别提到的情况，都可以认为符合正常使用条件。请只根据题目给出的情况和当前标签判断。你可以继续查看标签；如果不能确定，可以选择“不确定”。"), size=10.5, color=INK)

    page_break(doc)
    add_page_header(doc, "COUNTERBALANCING", "分组与顺序分配")
    rows = [list(r) for r in COUNTERBALANCE]
    add_simple_table(doc, ["组别", "Block 1", "Block 2", "Block 3"], rows, [25, 52, 52, 53], font_size=9.5)
    add_section_gap(doc, 8)
    add_notice(doc, "每位参与者依次完成3个 block；每个产品和每种标签条件各出现一次。分配登记表见下一页。", fill="F7F8F7", color="2C5F58")

    page_break(doc)
    add_page_header(doc, "ASSIGNMENT LOG", "参与者分配登记", note="建议按 G1-G6 循环分配；样本量超过18人时复印本页继续登记。")
    assignment_rows = []
    for i in range(1, 19):
        assignment_rows.append((str(i), "____________", f"G{((i - 1) % 6) + 1}", "____ / ____ / ____", ""))
    add_simple_table(doc, ["序号", "参与者编号", "组别", "日期", "备注"], assignment_rows, [18, 45, 24, 42, 53], font_size=7.8)

    page_break(doc)
    add_page_header(doc, "T1 · ANSWER KEY", "Product A 与 Product B 标准答案")
    add_notice(doc, "B3 已按标签事实修正为“本季已经喷过两次，还可以再喷一次”，答案为错；原题写“一次”会与每季最多2次相冲突。", fill="FFF8E9", color=GOLD)
    rows = []
    for product in ("A", "B"):
        for item, statement, answer, reason in T1[product]:
            unsafe = "是" if answer == "错" else "否"
            rows.append((item, statement, answer, reason, unsafe))
    add_simple_table(doc, ["题号", "陈述", "答案", "评分依据", "错判为“对”是否为不安全错误"], rows, [15, 70, 16, 58, 23], font_size=8.1)

    page_break(doc)
    add_page_header(doc, "T1 · ANSWER KEY", "Product C 标准答案与计分")
    rows = []
    for item, statement, answer, reason in T1["C"]:
        rows.append((item, statement, answer, reason, "是" if answer == "错" else "否"))
    add_simple_table(doc, ["题号", "陈述", "答案", "评分依据", "错判为“对”是否为不安全错误"], rows, [15, 70, 16, 58, 23], font_size=8.3)
    add_section_gap(doc, 7)
    p = doc.add_paragraph(style="Heading 2")
    set_run(p.add_run("T1 建议指标"), size=13, bold=True, color="2C5F58")
    metrics = [
        ("Judgment Accuracy", "每题答对记1，答错或不确定记0；也可单独报告不确定率。"),
        ("Unsafe-use Error", "正确答案为“错”却回答“对”时记1。"),
        ("Critical Information Miss", "根据错误题型编码为时期、间隔、次数、收获前、喷施方式或天气。"),
        ("Response Time / Confidence", "每页总时长；信心可用1-5分或本记录表中的三级选项。"),
    ]
    add_simple_table(doc, ["指标", "记录规则"], metrics, [48, 134], font_size=9)

    page_break(doc)
    add_page_header(doc, "T2 · ANSWER KEY", "混配卡标准分类")
    rows = []
    for product in ("A", "B", "C"):
        for card_id, name, ingredient, _, category in MIXING[product]:
            rows.append((f"Product {product}", card_id, name, ingredient, category))
    add_simple_table(doc, ["当前产品", "卡号", "现实产品", "卡片有效成分", "标准分类"], rows, [25, 28, 46, 48, 35], font_size=8.2)
    add_section_gap(doc, 5)
    add_notice(doc, "“标签未发现明确禁忌”不等于真实农业中一定可以混用。本任务只评价参与者能否依据给定实验标签完成分类。", fill="F7F8F7", color="2C5F58")
    p = doc.add_paragraph(style="Heading 2")
    set_run(p.add_run("关键错误编码"), size=13, bold=True, color="2C5F58")
    add_simple_table(doc, ["错误", "定义"], [
        ("False-Safe Error", "明确不能混用的卡被放入“标签未发现明确禁忌”。"),
        ("False-Prohibition Error", "未明确禁配的卡被放入“明确不能混用”。"),
        ("Uncertainty", "卡片被放入“不确定”。"),
    ], [52, 130], font_size=9)

    page_break(doc)
    add_page_header(doc, "T3 · GROUND TRUTH", "一亩地实际配药标准")
    add_notice(doc, "统一 ground truth：一亩总用水30 L。容器大小不作规定，10 L、15 L、20 L或其他容器均可；只评价总水量、按实际水量换算的药量和操作过程。", fill="EEF6F0", color=GREEN)
    rows = [
        ("A", "1.67 mL/L", "10 L=16.7；15 L=25；20 L=33.3", "50 mL/亩", "2.00 mL/L", "60 mL/亩"),
        ("B", "3.33 mL/L", "10 L=33.3；15 L=50；20 L=66.7", "100 mL/亩", "4.00 mL/L", "120 mL/亩"),
        ("C", "3.33 mL/L", "10 L=33.3；15 L=50；20 L=66.7", "100 mL/亩", "4.00 mL/L", "120 mL/亩"),
    ]
    add_simple_table(doc, ["产品", "一般浓度", "不同容器一般量（mL）", "一亩一般总量", "最高浓度", "一亩总上限"], rows, [16, 28, 66, 28, 22, 22], font_size=8.1)
    add_section_gap(doc, 7)
    p = doc.add_paragraph(style="Heading 2")
    set_run(p.add_run("正确操作顺序"), size=13, bold=True, color="2C5F58")
    add_simple_table(doc, ["步骤", "操作"], [
        ("1", "自行选择容器，并确定一桶或多桶的水量安排；所有桶合计30 L。"),
        ("2", "先加入约一半清水。"),
        ("3", "充分摇匀实验药瓶，按该桶实际水量换算并量取药剂。"),
        ("4", "加入药液并搅匀。"),
        ("5", "补水到该桶计划水量；如分多桶，逐桶按各自水量配制，合计30 L。"),
    ], [22, 160], font_size=8.8)
    p = doc.add_paragraph(style="Heading 2")
    set_run(p.add_run("计算规则"), size=13, bold=True, color="2C5F58")
    add_simple_table(doc, ["指标", "计算"], [
        ("Dose Deviation", "实际/计划一亩总药量 - 对应产品一亩一般总量。"),
        ("Upper-Bound Violation", "任一桶的药水比例超过最高浓度，或一亩总量超过总上限时记1。"),
        ("Procedure Error", "漏掉半水、按实际水量换算、量药、混匀、补足或多桶说明中的任一步骤。"),
    ], [52, 130], font_size=9)

    page_break(doc)
    add_page_header(doc, "T4 · CODING RUBRIC", "情境调整评分与产品标准")
    rows = [
        ("2 · Context-appropriate", "注意到附加条件；明确调整用量或覆盖方式；不突破标签上限。"),
        ("1 · Safe but non-adaptive", "维持一般方案，未出现危险操作，但没有根据新情境调整。"),
        ("0 · Unsafe / inappropriate", "任意加倍、超过上限、忽略关键边界，或提出明显不当方案。"),
    ]
    add_simple_table(doc, ["评分", "编码原则"], rows, [54, 128], font_size=9.2)
    add_section_gap(doc, 7)
    product_rows = [
        ("A", "杂草更密但仍矮小", "一般1.67 mL/L", "明确提高但不超过2.00 mL/L；一亩不超过60 mL，并强调对杂草茎叶喷匀。"),
        ("B", "上次效果明显下降", "一般3.33 mL/L", "明确提高但不超过4.00 mL/L；标签情境提示约3.67 mL/L，并说明不任意加倍。"),
        ("C", "低龄幼虫更多且上次效果不佳", "一般3.33 mL/L", "明确提高但不超过4.00 mL/L；标签情境提示约3.67 mL/L，并加强心叶/受害部位覆盖。"),
    ]
    add_simple_table(doc, ["产品", "变化", "一般方案", "2分参考"], product_rows, [18, 42, 35, 87], font_size=8.5)
    add_section_gap(doc, 6)
    add_notice(doc, "2分不要求逐字说出标签措辞；只要参与者明确响应情境、方案仍在边界内，并能说明理由即可。", fill="F7F8F7", color="2C5F58")

    page_break(doc)
    add_page_header(doc, "OBSERVATION FORM · 1/2", "Block 记录表：T1 与 T2")
    add_info_strip(doc, [("参与者编号", "____________"), ("组别 / Block", "G____ / ____"), ("产品 / 条件", "____ / O·E·D"), ("日期", "____年__月__日")], [48, 42, 48, 44])
    add_section_gap(doc, 6)
    p = doc.add_paragraph(style="Heading 2")
    set_run(p.add_run("T1 · 单变量判断"), size=13, bold=True, color="2C5F58")
    t1_blank = []
    for i in range(1, 6):
        t1_blank.append((str(i), "□ 对  □ 错  □ 不确定", "□ 0  □ 1", "□ 0  □ 1", "____"))
    add_simple_table(doc, ["题", "参与者答案", "正确", "不安全错误", "备注/标签回看"], t1_blank, [16, 55, 28, 35, 48], font_size=8.8)
    add_section_gap(doc, 4)
    add_info_strip(doc, [("开始", "____:____"), ("结束", "____:____"), ("总时长", "______秒"), ("信心", "□高 □中 □低")], [45.5] * 4)
    p = doc.add_paragraph(style="Heading 2")
    set_run(p.add_run("T2 · 混配分类"), size=13, bold=True, color="2C5F58")
    t2_blank = []
    for i in range(1, 5):
        t2_blank.append((str(i), "________________", "________________", "□ 0  □ 1", "□ FS  □ FP  □ U"))
    add_simple_table(doc, ["卡", "卡号/名称", "最终分类", "正确", "错误编码"], t2_blank, [14, 48, 58, 25, 37], font_size=8.8)
    add_section_gap(doc, 4)
    add_info_strip(doc, [("开始", "____:____"), ("结束", "____:____"), ("标签回看", "______次"), ("求助", "______次")], [45.5] * 4)
    p = doc.add_paragraph()
    set_paragraph(p, before=7, after=0)
    set_run(p.add_run("行为观察：________________________________________________________________________________\n__________________________________________________________________________________________"), size=9, color=INK)

    page_break(doc)
    add_page_header(doc, "OBSERVATION FORM · 2/2", "Block 记录表：T3 与 T4")
    add_info_strip(doc, [("参与者编号", "____________"), ("组别 / Block", "G____ / ____"), ("产品 / 条件", "____ / O·E·D"), ("日期", "____年__月__日")], [48, 42, 48, 44])
    p = doc.add_paragraph(style="Heading 2")
    set_run(p.add_run("T3 · 实际配药"), size=13, bold=True, color="2C5F58")
    add_simple_table(doc, ["项目", "记录"], [
        ("容器与桶数（不评分）", "选择：□10 L  □15 L  □20 L  □其他____；计划____桶"),
        ("总水量", "计划______L；最终/口述______L"),
        ("实际药量", "各桶：1____  2____  3____  4____ mL；一亩总量______mL"),
        ("量具", "□瓶盖  □纸杯  □量筒  □原瓶  □其他________"),
        ("操作", "□先半水  □摇匀药瓶  □按水量换算  □正确量药  □混匀  □补足/说明多桶"),
        ("计算", "Dose Deviation ______mL；Upper-Bound Violation □0 □1"),
    ], [48, 134], font_size=8.8)
    add_info_strip(doc, [("开始", "____:____"), ("结束", "____:____"), ("标签回看", "______次"), ("求助/计算", "______次")], [45.5] * 4)
    p = doc.add_paragraph(style="Heading 2")
    set_run(p.add_run("T4 · 情境调整"), size=13, bold=True, color="2C5F58")
    add_simple_table(doc, ["项目", "记录"], [
        ("是否继续使用", "□是  □否  □不确定"),
        ("最终方案", "最终浓度______mL/L；一亩总量______mL；调整方向 □增 □不变 □减"),
        ("解释", "________________________________________________________________"),
        ("编码", "T4得分 □0 □1 □2；Unsafe Escalation □0 □1"),
        ("过程", "开始____:____；结束____:____；信心 □高 □中 □低"),
    ], [48, 134], font_size=8.8)
    p = doc.add_paragraph()
    set_paragraph(p, before=7, after=0)
    set_run(p.add_run("补充观察：________________________________________________________________________________\n__________________________________________________________________________________________"), size=9, color=INK)

    doc.save(RESEARCHER_DOCX)


if __name__ == "__main__":
    build_participant_doc()
    build_researcher_doc()
    print(PARTICIPANT_DOCX)
    print(RESEARCHER_DOCX)
